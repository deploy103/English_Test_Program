import json
import random
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import shutil
import csv
import os
import subprocess
from docx import Document  # pip install python-docx

# ====== 데이터 저장 위치 (단어장/설정) ======
# exe만 배포 가능: 단어장은 사용자 AppData(Local) 아래에 저장/사용
APPDATA_DIR = Path(os.getenv("LOCALAPPDATA") or Path.home() / "AppData/Local")
APP_ROOT = APPDATA_DIR / "WordTest"
WORD_DIR = APP_ROOT / "word"           # 단어장 JSON 저장소
WORD_DIR.mkdir(parents=True, exist_ok=True)

# 테스트 횟수 저장 파일 (AppData)
TEST_COUNT_FILE = APP_ROOT / "test_count.txt"
TEST_COUNT_FILE.parent.mkdir(parents=True, exist_ok=True)

# 정답 저장 폴더(사용자가 지정, 강제)
SAVE_DIR: Path | None = None

# ====== 테스트 횟수 관리 ======
def load_test_count() -> int:
    if TEST_COUNT_FILE.exists():
        try:
            return int(TEST_COUNT_FILE.read_text(encoding="utf-8").strip())
        except:
            return 0
    return 0

def save_test_count(n: int) -> None:
    TEST_COUNT_FILE.write_text(str(n), encoding="utf-8")

test_count = load_test_count()

def reset_test_count():
    global test_count
    test_count = 0
    save_test_count(test_count)
    messagebox.showinfo("초기화", "테스트 횟수가 0으로 초기화되었습니다.")

# ====== (선택) 최초 실행 시 샘플 단어장 한 개 생성 ======
EMBEDDED_SAMPLE = [
    {"english": "astonished", "korean": "깜짝 놀란"},
    {"english": "immediately", "korean": "즉시"},
    {"english": "undamaged", "korean": "손상되지 않은"},
    {"english": "embarrassed", "korean": "당황스러운"},
    {"english": "valid", "korean": "타당한, 유효한"},
]
def seed_sample_if_empty():
    if not any(WORD_DIR.glob("*.json")):
        (WORD_DIR / "Sample_words.json").write_text(
            json.dumps(EMBEDDED_SAMPLE, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

seed_sample_if_empty()

# ====== 단어장 로드 ======
def load_words(path: Path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return [{k.lower(): v for k, v in w.items()} for w in data if "english" in w and "korean" in w]
    except Exception as e:
        messagebox.showerror("에러", f"파일 로드 실패:\n{e}")
        return []

# ====== 파일 추가 (AppData/word 폴더에 복사) ======
def add_file():
    filepath = filedialog.askopenfilename(
        title="단어장 JSON 파일 선택",
        filetypes=[("JSON files", "*.json")]
    )
    if not filepath:
        return
    src = Path(filepath)
    dest = WORD_DIR / src.name
    try:
        shutil.copy(src, dest)
    except Exception as e:
        messagebox.showerror("에러", f"파일 복사 실패:\n{e}")
        return
    update_file_list()
    messagebox.showinfo("완료", f"{src.name} 추가 완료!")

# ====== 파일명 → 별칭 변환 ======
def alias_name(filename: str) -> str:
    lower = filename.lower()
    alias = "[기타]"
    for n in range(1, 51):
        if f"lesson{n}" in lower:
            alias = f"[{n}단원]"
            break
    return f"{alias} {filename}"

# ====== 콤보박스 데이터 (별칭→실제파일명) ======
file_map: dict[str, str] = {}  # alias -> real filename

def update_file_list():
    global file_map
    file_map = {}
    for f in sorted(WORD_DIR.glob("*.json")):
        file_map[alias_name(f.name)] = f.name
    file_combo["values"] = list(file_map.keys())
    if file_map and not file_var.get():
        file_var.set(list(file_map.keys())[0])

# ====== 정답 저장 폴더 선택(강제) ======
def choose_save_dir():
    global SAVE_DIR
    path = filedialog.askdirectory(title="정답표 저장 폴더 선택")
    if path:
        SAVE_DIR = Path(path)
        messagebox.showinfo("저장 경로 선택됨", f"정답표는 여기 저장됩니다:\n{SAVE_DIR}")

# ====== 저장 함수들 (SAVE_DIR 필수) ======
def ensure_save_dir():
    if SAVE_DIR is None:
        raise RuntimeError("정답표 저장 폴더가 지정되지 않았습니다.")

def save_csv(history, test_number):
    ensure_save_dir()
    filename = SAVE_DIR / f"test_{test_number}.csv"
    with open(filename, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["번호", "문제", "정답"])
        for idx, q, a in history:
            writer.writerow([idx, q, a])
    return filename

def save_txt(history, test_number):
    ensure_save_dir()
    filename = SAVE_DIR / f"test_{test_number}.txt"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(f"테스트 {test_number} 결과\n")
        f.write("="*40 + "\n")
        f.write(f"{'번호':<6}{'문제':<20}{'정답':<20}\n")
        f.write("="*40 + "\n")
        for idx, q, a in history:
            f.write(f"{idx:<6}{q:<20}{a:<20}\n")
        f.write("="*40 + "\n")
    return filename

def save_docx(history, test_number):
    ensure_save_dir()
    filename = SAVE_DIR / f"test_{test_number}.docx"
    doc = Document()
    doc.add_heading(f"테스트 {test_number} 결과", 0)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "번호", "문제", "정답"
    for idx, q, a in history:
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(q)
        row[2].text = str(a)
    doc.save(filename)
    return filename

def open_folder(path: Path):
    try:
        if os.name == "nt":
            os.startfile(str(path))
        elif os.name == "posix":
            subprocess.call(["open", str(path)])
    except Exception as e:
        messagebox.showerror("에러", f"폴더 열기 실패:\n{e}")

# ====== 테스트 실행 ======
def start_test():
    # 저장 폴더 강제
    if SAVE_DIR is None:
        messagebox.showerror("에러", "테스트 시작 전, 정답표 저장 폴더를 먼저 지정하세요.")
        return

    alias = file_var.get()
    if not alias:
        messagebox.showwarning("경고", "단어장을 선택하세요.")
        return

    filename = file_map.get(alias)
    json_path = WORD_DIR / filename

    try:
        num_questions = int(count_var.get())
        display_time = int(time_var.get())
    except ValueError:
        messagebox.showwarning("경고", "숫자 값이 잘못되었습니다.")
        return

    show_mode = mode_var.get()

    words = load_words(json_path)
    if not words:
        return

    random.shuffle(words)
    selected = words[:min(num_questions, len(words))]

    history = []
    idx = 0

    # === 테스트 모드: 상단 UI 숨기고 단어만 크게 ===
    top_frame.pack_forget()
    btn_start.pack_forget()
    bottom_bar.pack_forget()
    label_word.config(text="3초 후 시작합니다...")

    def show_word():
        nonlocal idx
        if idx >= len(selected):
            finish_test()
            return

        w = selected[idx]
        en, ko = w["english"], w["korean"]

        if show_mode == "ko":
            shown, answer = ko, en
        elif show_mode == "en":
            shown, answer = en, ko
        else:
            shown = random.choice([en, ko])
            answer = ko if shown == en else en

        idx += 1
        label_word.config(text=f"{idx}. {shown}")
        history.append((idx, shown, answer))

        root.after(display_time * 1000, show_word)

    def finish_test():
        nonlocal history
        global test_count
        test_count += 1
        save_test_count(test_count)

        try:
            csv_file = save_csv(history, test_count)
            txt_file = save_txt(history, test_count)
            docx_file = save_docx(history, test_count)
        except Exception as e:
            messagebox.showerror("에러", f"정답 저장 실패:\n{e}")
            restore_home()
            return

        # === 결과창 ===
        result_win = tk.Toplevel(root)
        result_win.title(f"테스트 결과 (Test {test_count})")
        result_win.geometry("700x520")
        result_win.configure(bg="#ffffff")

        ttk.Label(result_win, text="✅ 테스트가 끝났습니다!",
                  font=("맑은 고딕", 16)).pack(pady=10)

        # 정답표 표시 영역
        text = tk.Text(result_win, width=80, height=22)
        text.pack(padx=12, pady=10, expand=True, fill="both")

        def show_answers():
            text.delete("1.0", "end")
            text.insert("end", f"테스트 {test_count} 결과\n\n")
            text.insert("end", f"{'번호':<6}{'문제':<28}{'정답':<28}\n")
            text.insert("end", "-"*70 + "\n")
            for i2, q, a in history:
                text.insert("end", f"{i2:<6}{q:<28}{a:<28}\n")
            text.insert("end", "-"*70 + "\n\n")
            text.insert("end", f"CSV: {csv_file}\nTXT: {txt_file}\nDOCX: {docx_file}\n")

        # 하단 버튼 바
        btn_bar = ttk.Frame(result_win)
        btn_bar.pack(side="bottom", fill="x", padx=10, pady=10)

        ttk.Button(btn_bar, text="정답표 보기", command=show_answers).pack(side="left", padx=5)
        ttk.Button(btn_bar, text="정답 파일 위치 열기",
                   command=lambda: open_folder(SAVE_DIR)).pack(side="left", padx=5)
        ttk.Button(btn_bar, text="홈으로 돌아가기",
                   command=lambda: (result_win.destroy(), restore_home())).pack(side="right", padx=5)

        label_word.config(text="테스트 완료. 결과창을 확인하세요.")

    # 카운트다운 → 시작
    root.after(1000, lambda: label_word.config(text="2..."))
    root.after(2000, lambda: label_word.config(text="1..."))
    root.after(3000, show_word)

def restore_home():
    label_word.config(text="테스트 대기 중...")
    top_frame.pack(fill="x", padx=12, pady=10)
    btn_start.pack(pady=(0, 12))
    bottom_bar.pack(fill="x", padx=12, pady=(0, 12))

# ====== GUI ======
root = tk.Tk()
root.title("Word Test Program")
root.geometry("820x600")
root.configure(bg="#ecf0f1")

style = ttk.Style()
style.theme_use("default")
style.configure("TButton", padding=6)
style.configure("TLabel", background="#ecf0f1")

# 상단 컨트롤 영역
top_frame = ttk.Frame(root, padding=10)
top_frame.pack(fill="x", padx=12, pady=10)

# 단어장 선택
ttk.Label(top_frame, text="단어장:").grid(row=0, column=0, sticky="w")
file_var = tk.StringVar()
file_combo = ttk.Combobox(top_frame, textvariable=file_var, width=46, state="readonly")
file_combo.grid(row=0, column=1, pady=4, sticky="w")

btn_add = ttk.Button(top_frame, text="파일 추가", command=add_file)
btn_add.grid(row=0, column=2, padx=6)

# 저장 폴더 지정 (강제)
btn_dir = ttk.Button(top_frame, text="정답표 저장 폴더 지정", command=choose_save_dir)
btn_dir.grid(row=0, column=3, padx=6)

# 문제 개수
ttk.Label(top_frame, text="문제 개수:").grid(row=1, column=0, sticky="w")
count_var = tk.StringVar(value="30")
ttk.Combobox(top_frame, textvariable=count_var, state="readonly",
             values=["30", "40", "50"], width=8).grid(row=1, column=1, sticky="w", pady=4)

# 언어 모드
ttk.Label(top_frame, text="언어 모드:").grid(row=1, column=2, sticky="e")
mode_var = tk.StringVar(value="rand")
ttk.Combobox(top_frame, textvariable=mode_var, state="readonly",
             values=["ko", "en", "rand"], width=8).grid(row=1, column=3, sticky="w", pady=4)

# 표시 시간
ttk.Label(top_frame, text="표시 시간(초):").grid(row=2, column=0, sticky="w")
time_var = tk.StringVar(value="3")
ttk.Combobox(top_frame, textvariable=time_var, state="readonly",
             values=["3", "5", "6"], width=8).grid(row=2, column=1, sticky="w", pady=4)

# 시작 버튼
btn_start = ttk.Button(root, text="테스트 시작", command=start_test)
btn_start.pack(pady=(0, 12))

# 하단 바(카운트 초기화)
bottom_bar = ttk.Frame(root, padding=10)
bottom_bar.pack(fill="x", padx=12, pady=(0, 12))
ttk.Button(bottom_bar, text="카운트 초기화", command=reset_test_count).pack(side="right")

# 단어 표시 영역 (테스트 모드에선 이것만 보임)
label_word = tk.Label(
    root, text="테스트 대기 중...",
    font=("맑은 고딕", 30, "bold"),
    fg="#2c3e50", bg="#ecf0f1",
    pady=30
)
label_word.pack(expand=True, fill="both")

# 초기 파일 목록
update_file_list()

root.mainloop()
